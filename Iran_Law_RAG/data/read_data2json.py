import docx
import re
import json


def clean_text(text: str) -> str:
    invisible_chars = ["\u200f", "\u200e", "\u200c\u200c", "\xa0"]
    for ch in invisible_chars:
        text = text.replace(ch, " " if ch == "\xa0" else "")
        text = text.replace( "\u200c", " ")
        text = text.replace(":", "")
    return re.sub(r"\s+", " ", text).strip()


def get_subject(paragraph, split_word):
    subj=[]
    for run in paragraph.runs:
        cl_text = clean_text(run.text)
        if run.bold is True and cl_text.startswith(split_word):
                subj.append(cl_text)

    return subj

def write_json(list_main_rules):
    json_main = json.dumps(list_main_rules, ensure_ascii=False, indent=4)
    with open('main_rule.json', 'w', encoding='utf-8') as f:
        f.write(json_main)
        

        
doc = docx.Document('./data/main_rule.docx')
print(doc)
asl = []
chapter = []
for paragraph in doc.paragraphs:
    text = paragraph.text
    if text:
        asl += get_subject(paragraph, "اصل")
        chapter += get_subject(paragraph, "فصل")

asl_specefic=''
chapter_specefic = ''
list_main_rules = []
temp_text = []
for paragraph in doc.paragraphs:
    lines = paragraph.text.splitlines()
    for line in lines:
         line = clean_text(line)
         if  line in chapter:
             chapter_specefic = line
             continue
         if line in asl:
              list_main_rules.append({
                        "chapter": chapter_specefic,
                        "asl": asl_specefic,
                        "text": temp_text.copy()
                    })

              temp_text = []
              asl_specefic=line
         else:
            temp_text.append(line)
print(list_main_rules)

write_json(list_main_rules)